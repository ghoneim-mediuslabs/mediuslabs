#!/usr/bin/env python3
"""
Markdown to DOCX Converter with HTML-to-Image Support

Converts markdown files to DOCX format, rendering embedded HTML blocks as images.
Supports both Arabic (RTL) and English (LTR) documents.

Usage:
    python md_to_docx.py <input.md> [output.docx]
    python md_to_docx.py docs/school-meals/prd-detailed.md
    python md_to_docx.py docs/school-meals/prd-detailed-ar.md --rtl
"""

import argparse
import asyncio
import hashlib
import os
import re
import sys
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def detect_language(content: str) -> str:
    """Detect if content is primarily Arabic or English."""
    arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]+')
    arabic_chars = len(arabic_pattern.findall(content))
    total_chars = len(re.findall(r'\w+', content))

    if total_chars == 0:
        return 'en'

    # If more than 30% Arabic characters, treat as Arabic
    return 'ar' if arabic_chars / max(total_chars, 1) > 0.3 else 'en'


def extract_html_blocks(content: str) -> list[tuple[str, int, int]]:
    """
    Extract HTML blocks from markdown content.
    Returns list of (html_content, start_pos, end_pos).
    """
    html_blocks = []

    # Pattern to match HTML blocks (div, table with styles, etc.)
    # Matches <div ...>...</div> blocks that span multiple lines
    pattern = re.compile(
        r'(<div\s+style=["\'][^"\']*["\'][^>]*>.*?</div>)',
        re.DOTALL | re.IGNORECASE
    )

    for match in pattern.finditer(content):
        html_content = match.group(1)
        # Only capture substantial HTML blocks (not simple inline elements)
        if len(html_content) > 100 and 'style=' in html_content:
            html_blocks.append((html_content, match.start(), match.end()))

    return html_blocks


async def html_to_image(html_content: str, output_path: str, is_rtl: bool = False) -> bool:
    """
    Render HTML content to an image using Playwright.
    """
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        print("Error: playwright not installed. Run: pip install playwright && playwright install chromium")
        return False

    # Wrap HTML in a full document with proper styling
    direction = 'rtl' if is_rtl else 'ltr'
    font_family = "'Segoe UI', 'Arial', 'Tahoma', sans-serif"

    full_html = f"""
    <!DOCTYPE html>
    <html dir="{direction}">
    <head>
        <meta charset="UTF-8">
        <style>
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            body {{
                font-family: {font_family};
                background: white;
                padding: 20px;
                direction: {direction};
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()

        await page.set_content(full_html)

        # Wait for any fonts/styles to load
        await page.wait_for_timeout(500)

        # Get the actual content size
        content_box = await page.evaluate('''() => {
            const body = document.body;
            const rect = body.getBoundingClientRect();
            return {
                width: Math.ceil(rect.width) + 40,
                height: Math.ceil(rect.height) + 40
            };
        }''')

        # Set viewport to content size
        await page.set_viewport_size({
            'width': min(content_box['width'], 1200),
            'height': content_box['height']
        })

        # Take screenshot
        await page.screenshot(path=output_path, full_page=True)
        await browser.close()

    return True


def create_html_template_for_table(table_md: str, is_rtl: bool = False) -> str:
    """Convert markdown table to styled HTML."""
    lines = table_md.strip().split('\n')
    if len(lines) < 2:
        return table_md

    direction = 'rtl' if is_rtl else 'ltr'
    align = 'right' if is_rtl else 'left'

    html = f'<table style="width: 100%; border-collapse: collapse; direction: {direction};">'

    # Header row
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    html += '<tr>'
    for header in headers:
        html += f'<th style="padding: 12px; background: #4CAF50; color: white; text-align: {align}; border: 1px solid #ddd;">{header}</th>'
    html += '</tr>'

    # Data rows (skip separator line)
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            html += '<tr>'
            for cell in cells:
                html += f'<td style="padding: 10px; text-align: {align}; border: 1px solid #ddd;">{cell}</td>'
            html += '</tr>'

    html += '</table>'
    return html


def parse_markdown_content(content: str) -> list[dict]:
    """
    Parse markdown content into structured blocks.
    Returns list of blocks with type and content.
    """
    blocks = []
    lines = content.split('\n')
    current_block = {'type': 'paragraph', 'content': ''}
    in_table = False
    table_content = []
    in_code_block = False
    code_content = []
    code_lang = ''

    i = 0
    while i < len(lines):
        line = lines[i]

        # Check for code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                if current_block['content'].strip():
                    blocks.append(current_block)
                    current_block = {'type': 'paragraph', 'content': ''}
                in_code_block = True
                code_lang = line.strip()[3:]
                code_content = []
            else:
                blocks.append({
                    'type': 'code',
                    'content': '\n'.join(code_content),
                    'language': code_lang
                })
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Check for HTML blocks
        if '<div' in line.lower() and 'style=' in line.lower():
            if current_block['content'].strip():
                blocks.append(current_block)
                current_block = {'type': 'paragraph', 'content': ''}

            # Collect the full HTML block
            html_content = line
            depth = line.count('<div') - line.count('</div>')

            while depth > 0 and i + 1 < len(lines):
                i += 1
                html_content += '\n' + lines[i]
                depth += lines[i].count('<div') - lines[i].count('</div>')

            blocks.append({'type': 'html', 'content': html_content})
            i += 1
            continue

        # Check for headings
        if line.startswith('#'):
            if current_block['content'].strip():
                blocks.append(current_block)
                current_block = {'type': 'paragraph', 'content': ''}

            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            blocks.append({'type': 'heading', 'content': text, 'level': level})
            i += 1
            continue

        # Check for tables
        if '|' in line and not line.strip().startswith('<'):
            if not in_table:
                if current_block['content'].strip():
                    blocks.append(current_block)
                    current_block = {'type': 'paragraph', 'content': ''}
                in_table = True
                table_content = []
            table_content.append(line)

            # Check if next line is not a table
            if i + 1 >= len(lines) or '|' not in lines[i + 1]:
                blocks.append({'type': 'table', 'content': '\n'.join(table_content)})
                in_table = False
                table_content = []
            i += 1
            continue

        # Check for horizontal rules
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            if current_block['content'].strip():
                blocks.append(current_block)
                current_block = {'type': 'paragraph', 'content': ''}
            blocks.append({'type': 'hr', 'content': ''})
            i += 1
            continue

        # Check for list items
        if re.match(r'^[\s]*[-*+]\s', line) or re.match(r'^[\s]*\d+\.\s', line):
            if current_block['type'] != 'list':
                if current_block['content'].strip():
                    blocks.append(current_block)
                current_block = {'type': 'list', 'content': line + '\n'}
            else:
                current_block['content'] += line + '\n'
            i += 1
            continue

        # Regular paragraph content
        if current_block['type'] == 'list' and line.strip():
            blocks.append(current_block)
            current_block = {'type': 'paragraph', 'content': ''}

        current_block['content'] += line + '\n'
        i += 1

    # Add remaining content
    if current_block['content'].strip():
        blocks.append(current_block)

    return blocks


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL direction."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def process_inline_formatting(paragraph, text: str, is_rtl: bool = False):
    """Process inline markdown formatting (bold, italic, code, links)."""
    # Pattern for various inline elements
    patterns = [
        (r'\*\*(.+?)\*\*', 'bold'),
        (r'__(.+?)__', 'bold'),
        (r'\*(.+?)\*', 'italic'),
        (r'_(.+?)_', 'italic'),
        (r'`(.+?)`', 'code'),
        (r'\[([^\]]+)\]\(([^)]+)\)', 'link'),
    ]

    # Simple approach: just add text with basic formatting detection
    remaining = text

    while remaining:
        # Find the earliest match
        earliest_match = None
        earliest_pos = len(remaining)
        match_type = None

        for pattern, fmt_type in patterns:
            match = re.search(pattern, remaining)
            if match and match.start() < earliest_pos:
                earliest_match = match
                earliest_pos = match.start()
                match_type = fmt_type

        if earliest_match is None:
            # No more formatting, add remaining text
            if remaining.strip():
                run = paragraph.add_run(remaining)
                if is_rtl:
                    run.font.rtl = True
            break

        # Add text before the match
        if earliest_pos > 0:
            run = paragraph.add_run(remaining[:earliest_pos])
            if is_rtl:
                run.font.rtl = True

        # Add formatted text
        if match_type == 'bold':
            run = paragraph.add_run(earliest_match.group(1))
            run.bold = True
        elif match_type == 'italic':
            run = paragraph.add_run(earliest_match.group(1))
            run.italic = True
        elif match_type == 'code':
            run = paragraph.add_run(earliest_match.group(1))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif match_type == 'link':
            link_text = earliest_match.group(1)
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True

        if is_rtl and match_type != 'code':
            run.font.rtl = True

        remaining = remaining[earliest_match.end():]


def create_table_in_docx(doc, table_md: str, is_rtl: bool = False):
    """Create a formatted table in the DOCX document."""
    lines = [l for l in table_md.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return

    # Parse headers
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    if not headers:
        return

    # Count columns
    num_cols = len(headers)

    # Create table
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'

    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        # Style header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_rtl else WD_ALIGN_PARAGRAPH.LEFT
            if is_rtl:
                set_rtl_paragraph(paragraph)
            for run in paragraph.runs:
                run.bold = True

    # Add data rows (skip separator line at index 1)
    for line in lines[2:]:
        cells_data = [cell.strip() for cell in line.split('|') if cell.strip()]
        if not cells_data:
            continue

        row = table.add_row()
        for i, cell_text in enumerate(cells_data[:num_cols]):
            cell = row.cells[i]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_rtl else WD_ALIGN_PARAGRAPH.LEFT
                if is_rtl:
                    set_rtl_paragraph(paragraph)

    # Add spacing after table
    doc.add_paragraph()


async def convert_md_to_docx(
    input_path: str,
    output_path: Optional[str] = None,
    force_rtl: bool = False,
    keep_images: bool = False
) -> str:
    """
    Convert a markdown file to DOCX format.

    Args:
        input_path: Path to the input markdown file
        output_path: Path for output DOCX (defaults to same name with .docx extension)
        force_rtl: Force RTL direction (for Arabic documents)
        keep_images: Keep generated images after conversion

    Returns:
        Path to the generated DOCX file
    """
    input_path = Path(input_path).resolve()

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(output_path).resolve()

    # Read input file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Detect language if not forced
    is_rtl = force_rtl or detect_language(content) == 'ar'

    print(f"Processing: {input_path}")
    print(f"Direction: {'RTL (Arabic)' if is_rtl else 'LTR (English)'}")

    # Parse content into blocks
    blocks = parse_markdown_content(content)

    # Create temporary directory for images
    temp_dir = tempfile.mkdtemp(prefix='md2docx_')
    image_paths = []

    # Create DOCX document
    doc = Document()

    # Set document direction for RTL
    if is_rtl:
        sections = doc.sections
        for section in sections:
            sectPr = section._sectPr
            bidi = OxmlElement('w:bidi')
            sectPr.append(bidi)

    # Process each block
    image_counter = 0
    for block in blocks:
        block_type = block['type']
        block_content = block['content']

        if block_type == 'heading':
            level = block.get('level', 1)
            style_name = f'Heading {min(level, 9)}'
            para = doc.add_paragraph(block_content, style=style_name)
            if is_rtl:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_rtl_paragraph(para)

        elif block_type == 'html':
            # Convert HTML to image
            image_counter += 1
            image_hash = hashlib.md5(block_content.encode()).hexdigest()[:8]
            image_path = os.path.join(temp_dir, f'html_block_{image_counter}_{image_hash}.png')

            print(f"  Rendering HTML block {image_counter} to image...")
            success = await html_to_image(block_content, image_path, is_rtl)

            if success and os.path.exists(image_path):
                image_paths.append(image_path)
                # Add image to document
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_path, width=Inches(6))
            else:
                # Fallback: add as text note
                para = doc.add_paragraph("[HTML content - see original document]")
                para.italic = True

        elif block_type == 'table':
            create_table_in_docx(doc, block_content, is_rtl)

        elif block_type == 'list':
            for line in block_content.strip().split('\n'):
                # Detect list type
                if re.match(r'^\s*\d+\.', line):
                    # Numbered list
                    text = re.sub(r'^\s*\d+\.\s*', '', line)
                    para = doc.add_paragraph(style='List Number')
                else:
                    # Bullet list
                    text = re.sub(r'^\s*[-*+]\s*', '', line)
                    para = doc.add_paragraph(style='List Bullet')

                process_inline_formatting(para, text, is_rtl)
                if is_rtl:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    set_rtl_paragraph(para)

        elif block_type == 'code':
            para = doc.add_paragraph()
            run = para.add_run(block_content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        elif block_type == 'hr':
            # Add horizontal line
            para = doc.add_paragraph('_' * 50)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif block_type == 'paragraph':
            text = block_content.strip()
            if text:
                para = doc.add_paragraph()
                process_inline_formatting(para, text, is_rtl)
                if is_rtl:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    set_rtl_paragraph(para)

    # Save document
    doc.save(str(output_path))
    print(f"Saved: {output_path}")

    # Cleanup temporary images
    if not keep_images:
        for img_path in image_paths:
            try:
                os.remove(img_path)
            except:
                pass
        try:
            os.rmdir(temp_dir)
        except:
            pass
    else:
        print(f"Images saved in: {temp_dir}")

    return str(output_path)


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to DOCX with HTML-to-image support',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python md_to_docx.py document.md
    python md_to_docx.py document.md output.docx
    python md_to_docx.py document-ar.md --rtl
    python md_to_docx.py document.md --keep-images
        """
    )

    parser.add_argument('input', help='Input markdown file path')
    parser.add_argument('output', nargs='?', help='Output DOCX file path (optional)')
    parser.add_argument('--rtl', action='store_true', help='Force RTL direction (Arabic)')
    parser.add_argument('--ltr', action='store_true', help='Force LTR direction (English)')
    parser.add_argument('--keep-images', action='store_true', help='Keep generated images')

    args = parser.parse_args()

    # Determine RTL setting
    force_rtl = args.rtl
    if args.ltr:
        force_rtl = False

    try:
        result = asyncio.run(convert_md_to_docx(
            args.input,
            args.output,
            force_rtl=force_rtl,
            keep_images=args.keep_images
        ))
        print(f"\nConversion complete: {result}")
    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error during conversion: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
